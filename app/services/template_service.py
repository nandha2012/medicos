from docx import Document
from services.pdf_service import PDFService
import os
from datetime import datetime

output_dir = os.getenv("OUTPUT_DIR") or "output"
def generate_dir_name():
    now = datetime.now()
    return now.strftime("%Y%m%d%H")

class TemplateService:
    
    def __init__(self, template_path: str):
        self.template_path = template_path
        self.output_path_docx = None
        self.output_dir = output_dir+"/"+generate_dir_name()
        os.makedirs(self.output_dir, exist_ok=True)
    def merge_runs(self, paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        for run in paragraph.runs:
            run.text = ''
        paragraph.runs[0].text = full_text

    def _replace_placeholders_in_paragraph(self, para, data):
        text = para.text
        for key, value in data.items():
            text = text.replace(f"#{key}#", str(value))

        # Remove all runs cleanly
        for i in reversed(range(len(para.runs))):
            para._element.remove(para.runs[i]._element)

        # Add cleaned run
        para.add_run(text)


    def fill_template(self, output_path: str, data: dict,j):
        print(f"📄 Using template: {self.template_path}")
        mg_idpreg = data.get('mg_idpreg')
        if mg_idpreg is None:
            raise KeyError("❌ 'mg_idpreg' key not found in data dictionary")
        file_path = self.output_dir+"/"+output_path
        os.makedirs(file_path, exist_ok=True)
        self.output_path_docx = os.path.join(file_path, f"{mg_idpreg}_{j}.docx")
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"❌ Template not found at: {self.template_path}")
        doc = Document(self.template_path)
        # Replace in paragraphs (runs preserve formatting)
        for para in doc.paragraphs:
            for run in para.runs:
                for key, value in data.items():
                    if f"#{key}#" in run.text:
                        print(f'replaceing #{key}# with {value}')
                        run.text = run.text.replace(f"#{key}#", str(value))

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for key, value in data.items():
                                if f"#{key}#" in run.text:
                                    run.text = run.text.replace(f"#{key}#", str(value))

        doc.save(self.output_path_docx)
        return self.output_path_docx
